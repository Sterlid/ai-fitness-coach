from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AI_Fitness_Coach_PRD_v0.1.docx"
BLUE = "2563EB"; NAVY = "17324D"; PALE = "EAF2FF"; GRAY = "5F6B7A"; LIGHT = "F3F5F7"; WHITE = "FFFFFF"; BLACK = "111827"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1); sec.right_margin = Inches(1)
sec.header_distance = Inches(0.492); sec.footer_distance = Inches(0.492)

def font(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [("Title",25,NAVY,0,5),("Subtitle",12,GRAY,0,14),("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,NAVY,8,4)]:
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=name not in ("Subtitle",)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

if "Callout" not in styles:
    s=styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH); s.base_style=styles["Normal"]
else: s=styles["Callout"]
s.paragraph_format.left_indent=Inches(.18); s.paragraph_format.right_indent=Inches(.18); s.paragraph_format.space_before=Pt(5); s.paragraph_format.space_after=Pt(9)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def table(headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.autofit=False; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; shade(c,NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),9,bold=True,color=WHITE)
    trPr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); trPr.append(repeat)
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        trPr=t.rows[-1]._tr.get_or_add_trPr(); cant=OxmlElement('w:cantSplit'); trPr.append(cant)
        for i,val in enumerate(row):
            c=cells[i]; c.text=""; set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx%2: shade(c,"F8FAFC")
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(val)),9,color=BLACK)
    if widths:
        width_dxa = [round(w * 1440) for w in widths]
        tbl_pr = t._tbl.tblPr
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None: tbl_w = OxmlElement('w:tblW'); tbl_pr.append(tbl_w)
        tbl_w.set(qn('w:w'), str(sum(width_dxa))); tbl_w.set(qn('w:type'), 'dxa')
        tbl_ind = tbl_pr.find(qn('w:tblInd'))
        if tbl_ind is None: tbl_ind = OxmlElement('w:tblInd'); tbl_pr.append(tbl_ind)
        tbl_ind.set(qn('w:w'), '120'); tbl_ind.set(qn('w:type'), 'dxa')
        layout = tbl_pr.find(qn('w:tblLayout'))
        if layout is None: layout = OxmlElement('w:tblLayout'); tbl_pr.append(layout)
        layout.set(qn('w:type'), 'fixed')
        grid_cols = t._tbl.tblGrid.findall(qn('w:gridCol'))
        for i, value in enumerate(width_dxa):
            grid_cols[i].set(qn('w:w'), str(value))
        for row in t.rows:
            for i,value in enumerate(width_dxa):
                row.cells[i].width=Inches(widths[i])
                tc_w = row.cells[i]._tc.get_or_add_tcPr().find(qn('w:tcW'))
                tc_w.set(qn('w:w'), str(value)); tc_w.set(qn('w:type'), 'dxa')
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def bullet(text, level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2')
    p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25)
    p.paragraph_format.space_after=Pt(8); p.paragraph_format.line_spacing=1.167; p.add_run(text); return p

def numbered(text):
    p=doc.add_paragraph(style='List Number'); p.paragraph_format.left_indent=Inches(.5)
    p.paragraph_format.first_line_indent=Inches(-.25); p.paragraph_format.space_after=Pt(8)
    p.paragraph_format.line_spacing=1.167; p.add_run(text); return p

def restart_numbering(paragraph, start=1):
    numbering = doc.part.numbering_part.element
    style_num_id = int(styles['List Number']._element.pPr.numPr.numId.val)
    base_num = next(n for n in numbering.findall(qn('w:num')) if int(n.get(qn('w:numId'))) == style_num_id)
    abstract_id = int(base_num.find(qn('w:abstractNumId')).get(qn('w:val')))
    ids = [int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num'))]
    new_id = max(ids) + 1
    num = OxmlElement('w:num'); num.set(qn('w:numId'), str(new_id))
    aid = OxmlElement('w:abstractNumId'); aid.set(qn('w:val'), str(abstract_id)); num.append(aid)
    override = OxmlElement('w:lvlOverride'); override.set(qn('w:ilvl'), '0')
    start_node = OxmlElement('w:startOverride'); start_node.set(qn('w:val'), str(start)); override.append(start_node); num.append(override)
    numbering.append(num)
    numPr = paragraph._p.get_or_add_pPr().get_or_add_numPr(); numPr.get_or_add_numId().val = new_id
    return new_id

def continue_numbering(paragraph, num_id):
    numPr = paragraph._p.get_or_add_pPr().get_or_add_numPr(); numPr.get_or_add_numId().val = num_id

def callout(label, text):
    p=doc.add_paragraph(style='Callout'); shade_dummy = OxmlElement('w:shd'); shade_dummy.set(qn('w:fill'),PALE); p._p.get_or_add_pPr().append(shade_dummy)
    font(p.add_run(label+"  "),10.5,bold=True,color=BLUE); font(p.add_run(text),10.5,color=BLACK)

# Header/footer
h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(h.add_run("PRODUCT REQUIREMENTS  •  V0.1"),8.5,bold=True,color=GRAY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(f.add_run("AI Fitness Coach  |  Working draft  |  August 2026"),8,color=GRAY)

# Cover/masthead
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(28); font(p.add_run("PRODUCT REQUIREMENTS DOCUMENT"),9,bold=True,color=BLUE)
p=doc.add_paragraph(style='Title'); p.add_run("AI Fitness Coach")
p=doc.add_paragraph(style='Subtitle'); p.add_run("Photo-first calorie tracking, adaptive meal guidance, and daily workout coaching")
table(["Version","Status","Owner","Last updated"], [["0.1","Discovery draft","Product","1 August 2026"]], [1.1,1.6,2.0,1.8])
callout("Product thesis", "People are more consistent when logging is nearly effortless and each log immediately improves a practical recommendation for what to eat or do next.")

doc.add_heading('1. Executive summary', level=1)
doc.add_paragraph("AI Fitness Coach is a mobile-first consumer app that turns meal photos or descriptions and lightweight activity data into a useful daily plan. It estimates calories and macronutrients, lets users correct uncertain details, recommends meals that fit the remainder of the day, and adapts workout suggestions to goals, history, recovery, time, and available equipment.")
doc.add_paragraph("The first release should optimize for trust and repeat use, not clinical precision. Estimates must show uncertainty, recommendations must explain their rationale, and users must retain control over edits, goals, dietary preferences, injuries, and data sharing.")

doc.add_heading('2. Problem and opportunity', level=1)
doc.add_heading('User problem', level=2)
for x in ["Manual food logging is slow, especially for mixed dishes and restaurant meals.","Calorie trackers often report numbers without helping users decide what to eat next.","Workout plans are frequently static and ignore adherence, fatigue, available time, or recent training.","Nutrition and exercise data live in separate experiences, so the daily plan does not adapt as behavior changes."]: bullet(x)
doc.add_heading('Opportunity', level=2)
doc.add_paragraph("Combine fast multimodal logging with an adaptive decision layer. Each meal or workout becomes both a record and a signal that updates the user's next-best action. The product wins when users feel understood, can correct mistakes quickly, and return because the guidance becomes more relevant over time.")

doc.add_heading('3. Goals and non-goals', level=1)
table(["Goals for MVP","Not in MVP"], [
    ("Log a meal from a photo or natural-language description in under 30 seconds.","Diagnose conditions or provide medical nutrition therapy."),
    ("Give editable calorie and macro estimates with a confidence cue.","Guarantee laboratory-grade nutrition accuracy."),
    ("Recommend meals based on remaining targets and preferences.","Grocery delivery, restaurant ordering, or payment."),
    ("Generate a safe daily workout and adapt it using recent behavior.","Live form analysis, wearables coaching, or social competition."),
    ("Create a clear daily view linking food, movement, and progress.","Support coaches, families, or enterprise programs."),
], [3.25,3.25])

doc.add_heading('4. Target users and jobs to be done', level=1)
table(["Segment","Primary job","Key need"], [
    ("Consistency seeker","Help me make better choices without obsessive tracking.","Low-friction logging and gentle guidance."),
    ("Goal-driven improver","Help me lose fat, maintain, or gain muscle predictably.","Targets, trends, and actionable adjustments."),
    ("Busy exerciser","Tell me what workout fits today.","Time-, recovery-, and equipment-aware plans."),
], [1.35,3.15,2.0])
doc.add_heading('Core jobs', level=2)
for x in ["When I eat, let me capture it quickly and understand the likely nutritional impact.","When I am deciding what to eat, show options that fit my remaining targets and constraints.","When I am ready to train, give me an achievable workout that complements my recent activity.","When estimates or plans are wrong, let me fix them easily so future guidance improves."]: bullet(x)

doc.add_heading('5. Product principles', level=1)
for label,text in [("Fast before perfect","A useful estimate now, followed by easy correction, beats a long intake form."),("Confidence is visible","Distinguish observed facts, inferred portions, and user-confirmed values."),("Recommendations explain themselves","Show the two or three factors that drove each suggestion."),("Adapt, do not punish","Missed logs and workouts should lead to a gentler next step, never shame."),("Safety and agency","Users control goals and constraints; the app avoids diagnosis and unsafe exercise escalation.")]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.keep_together=True
    font(p.add_run(label+": "),10.5,bold=True,color=NAVY); p.add_run(text)

doc.add_heading('6. MVP scope and experience', level=1)
doc.add_heading('6.1 Onboarding and profile', level=2)
for x in ["Collect age range, height, weight, activity level, primary goal, preferred pace, dietary preferences, allergies, disliked foods, training experience, available equipment, usual workout duration, limitations/injuries, and preferred units.","Calculate a suggested calorie range and macro targets, clearly labeled as estimates; require user confirmation.","Allow users to skip optional fields and revise all goals later."]: bullet(x)
doc.add_heading('6.2 Meal capture and analysis', level=2)
for x in ["Accept camera photo, library image, text description, or photo plus text.","Identify likely foods, preparation method, and portion ranges; ask at most two high-value clarification questions when uncertainty materially changes the estimate.","Return calories, protein, carbohydrates, fat, serving assumptions, and a confidence level.","Let the user edit items, portions, ingredients, and totals; save corrections as personalization signals.","Support saved meals, recent meals, copy-to-today, deletion, and meal time/category."]: bullet(x)
doc.add_heading('6.3 Meal recommendations', level=2)
for x in ["Recommend 3–5 options using remaining daily targets, dietary constraints, recent variety, preparation time, and user preferences.","Each recommendation includes a name, short rationale, approximate calories/macros, ingredients, and simple preparation steps.","Users can request alternatives, exclude an ingredient, change calorie range, save a meal, or mark not interested."]: bullet(x)
doc.add_heading('6.4 Daily workout coaching', level=2)
for x in ["Generate one primary workout plus a lower-intensity alternative and a rest/recovery option.","Use goal, experience, recent muscle-group exposure, completed workouts, reported soreness/energy, available time, equipment, and limitations.","Represent workouts as warm-up, exercises with sets/reps/rest or time, cooldown, estimated duration, and intensity cue.","Allow replace exercise, shorten workout, change equipment, mark set/workout complete, rate difficulty, report pain, or skip with a reason.","Never recommend training through pain; persistent or severe symptoms trigger a stop-and-seek-professional-care message."]: bullet(x)
doc.add_heading('6.5 Today and progress', level=2)
for x in ["Today screen shows calorie/macro progress, logged meals, next meal suggestions, planned/completed workout, and one prioritized next action.","Weekly view shows logging consistency, calorie and protein ranges, workout completion, training balance, and weight trend when supplied.","Avoid single-day judgment; frame insights around rolling patterns and user-defined goals."]: bullet(x)

doc.add_heading('7. Key user flows', level=1)
doc.add_heading('Meal photo to logged entry', level=2)
flow_num=None
for x in ["User photographs a meal and optionally adds context.","System detects meal components and portions, then flags important uncertainty.","User confirms or adjusts the estimate.","Meal is added to the day; remaining targets and recommendations refresh."]:
    p=numbered(x)
    if flow_num is None: flow_num=restart_numbering(p)
    else: continue_numbering(p,flow_num)
doc.add_heading('Daily workout to completion', level=2)
flow_num=None
for x in ["User opens Today and sees the recommended session plus rationale.","User confirms time/equipment and optionally reports energy or soreness.","System adapts the session; user logs completion and perceived difficulty.","Future plans update based on adherence, training load, and feedback."]:
    p=numbered(x)
    if flow_num is None: flow_num=restart_numbering(p)
    else: continue_numbering(p,flow_num)

doc.add_heading('8. Functional requirements', level=1)
table(["ID","Requirement","Priority"], [
    ("FR-01","Create and edit a personal profile, goals, food constraints, and training constraints.","P0"),
    ("FR-02","Analyze a meal image and/or text and produce itemized editable estimates.","P0"),
    ("FR-03","Store meal logs and update daily calorie/macro totals immediately.","P0"),
    ("FR-04","Generate personalized meal recommendations with rationale and alternatives.","P0"),
    ("FR-05","Generate a daily workout with safe adaptations and completion logging.","P0"),
    ("FR-06","Show today and weekly progress views.","P0"),
    ("FR-07","Capture thumbs up/down, edits, skips, difficulty, soreness, and pain feedback.","P0"),
    ("FR-08","Support reminders with explicit opt-in and granular controls.","P1"),
    ("FR-09","Integrate Apple Health / Health Connect for activity and workout import.","P1"),
    ("FR-10","Export and delete account data.","P0"),
], [.7,4.95,.85])

doc.add_heading('9. Recommendation logic', level=1)
doc.add_heading('Meal ranking inputs', level=2)
doc.add_paragraph("Hard filters: allergies, dietary exclusions, and explicit dislikes. Ranking factors: remaining calorie/protein targets, meal timing, recent food variety, saved preferences, preparation time, available ingredients (when known), and prior acceptance. Recommendations should not compensate for an over-target day with extreme restriction.")
doc.add_heading('Workout planning inputs', level=2)
doc.add_paragraph("Hard filters: injuries/limitations, equipment availability, experience-appropriate movements, and pain reports. Ranking and adaptation factors: user goal, time, recent sessions, muscle-group recovery, adherence, perceived difficulty, soreness, sleep/energy if voluntarily supplied, and progression history. Load or volume changes should be conservative and reversible.")
callout("Model behavior requirement", "The system must provide structured output, retain the source inputs used, expose uncertainty, and fall back to a safe generic option when confidence or data quality is low.")

doc.add_heading('10. Data and system model', level=1)
table(["Entity","Key fields"], [
    ("User profile","Goals, demographics, units, preferences, allergies, limitations, consent settings"),
    ("Meal entry","Timestamp, image reference, text, items, portions, nutrients, confidence, edits"),
    ("Meal recommendation","Candidate, rationale, constraints, predicted nutrients, response"),
    ("Workout plan","Date, goal, duration, equipment, exercises, dosage, adaptations, status"),
    ("Workout log","Completed work, difficulty, soreness/pain, substitutions, notes"),
    ("Daily target","Calorie/macro range, derivation version, user override"),
    ("Feedback event","Action, context, recommendation/version identifiers"),
], [1.55,4.95])
doc.add_paragraph("Images should be encrypted in transit and at rest, retained only as needed for the user-facing feature and stated improvement program, and removable by the user. Separate personally identifying data from model telemetry where practical. Maintain versioned estimates so changed model behavior does not silently rewrite history.")
doc.add_heading('Technical direction', level=2)
doc.add_paragraph("The initial client is a cross-platform Expo and React Native application written in TypeScript. Supabase provides authentication, Postgres, Row Level Security, and private object storage. AI-provider calls and privileged business logic must run server-side through Supabase Edge Functions; provider secrets and service-role credentials must never be embedded in the mobile client.")

doc.add_heading('11. Trust, safety, privacy, and accessibility', level=1)
for x in ["Require adult eligibility for MVP or implement an age-appropriate experience before supporting minors.","Use calorie ranges and uncertainty cues; do not present image-based estimation as exact.","Screen for high-risk use patterns and avoid aggressive deficits, punitive language, compensatory exercise, or recommendations associated with disordered eating.","For pregnancy, eating-disorder history, complex medical conditions, or significant injury, limit guidance and recommend qualified professional support.","Provide consent, data-access, export, deletion, notification, analytics, and model-improvement controls.","Meet WCAG 2.2 AA intent: screen-reader labels, dynamic type, contrast, non-color status cues, captions/alt text, and accessible camera alternatives.","Red-team image uploads, prompt injection, unsafe workout plans, allergen failures, and misleading health claims before launch."]: bullet(x)

doc.add_heading('12. Non-functional requirements', level=1)
table(["Area","MVP target"], [
    ("Performance","First analysis response in ≤8 seconds at p75; full result in ≤15 seconds at p90 under normal conditions."),
    ("Availability","99.5% monthly for core logging and Today APIs."),
    ("Reliability","Draft capture survives transient failure; retries do not create duplicate entries."),
    ("Security","Encryption in transit/at rest, least privilege, secret management, audit logging, dependency scanning."),
    ("Observability","Trace recommendation versions, latency, failure mode, user correction, and safety intervention."),
    ("Localization","Architecture supports locale, units, food vocabularies, and time zones; MVP language TBD."),
], [1.35,5.15])

doc.add_heading('13. Success metrics', level=1)
doc.add_heading('North-star behavior', level=2)
doc.add_paragraph("Weekly Guided Days: number of days per week on which an active user logs at least one meal and either views, adapts, or completes the recommended workout/meal action.")
table(["Metric","Initial hypothesis / guardrail"], [
    ("Activation","≥60% of new users complete onboarding and log a first meal within 24 hours."),
    ("Logging value","≥70% of analyzed meals are saved; median completion time <30 seconds."),
    ("Estimate trust","<35% of saved meals require a material calorie edit after the first four weeks."),
    ("Recommendation value","≥25% weekly users save, select, or positively rate a meal recommendation."),
    ("Workout value","≥35% weekly users complete at least two recommended workouts."),
    ("Retention","D30 retained-user target to be set after closed beta baseline."),
    ("Safety guardrail","Zero known severe preventable allergen or contraindicated-exercise incidents; investigate all reports."),
], [2.0,4.5])

doc.add_heading('14. Analytics and experimentation', level=1)
for x in ["Instrument onboarding completed, meal capture started, analysis returned/failed, clarification shown, edit made, meal saved, recommendation viewed/selected/dismissed, workout generated/adapted/completed/skipped, pain reported, and account/data controls used.","Attach model/version, confidence band, latency, and anonymized decision context to relevant events.","Do not optimize only for engagement; pair experiments with correction rate, opt-outs, safety signals, and self-reported usefulness.","Initial experiments: clarification threshold, confidence presentation, number of recommendations, Today screen priority card, and workout adaptation prompt."]: bullet(x)

doc.add_heading('15. Release plan and acceptance criteria', level=1)
table(["Phase","Exit criteria"], [
    ("Prototype","End-to-end clickable experience; 8–12 moderated tests validate comprehension and correction flow."),
    ("Internal alpha","Core flows functional; structured outputs validated; critical security and safety tests pass."),
    ("Closed beta","≥200 consented testers; quality dashboards live; support and incident process operational."),
    ("MVP launch","P0 requirements complete; store/privacy review passed; on-call, rollback, and model version controls ready."),
], [1.35,5.15])
doc.add_heading('MVP acceptance criteria', level=2)
for x in ["A user can complete onboarding, log and edit a meal, see updated targets, receive meal options, generate/adapt/complete a workout, and view weekly progress.","Every AI estimate or recommendation contains a confidence or rationale signal and can be corrected or dismissed.","Allergy and physical-limitation constraints are enforced in test suites and logged for audit.","The app remains usable when image analysis fails by offering text/manual entry.","A user can export and delete their data from settings."]: bullet(x)

doc.add_heading('16. Risks and mitigations', level=1)
table(["Risk","Mitigation"], [
    ("Inaccurate portions or ingredients","Show ranges and assumptions; ask targeted questions; prioritize easy edits."),
    ("Unsafe or overconfident health guidance","Conservative rules, safety classifiers, expert review, escalation language, audit logs."),
    ("User fixation or disordered patterns","Ranges, neutral language, configurable visibility, risk detection, restricted advice."),
    ("Cold-start recommendations","Short onboarding, generic safe defaults, rapid learning from explicit feedback."),
    ("Privacy concerns around meal images and health data","Data minimization, clear retention, user controls, encryption, no hidden reuse."),
    ("Recommendation fatigue","One prioritized next action, bounded notifications, diversity and dismissal learning."),
], [2.05,4.45])

doc.add_heading('17. Open decisions', level=1)
for x in ["Primary initial goal segment: general wellness, fat loss, or strength/muscle gain?","Market and language for launch, which determine nutrition databases and regulatory review.","Whether calorie targets are always shown or can be replaced by non-calorie habit goals.","Nutrition data provider and image-model architecture; cost, latency, and licensing constraints.","Authentication options beyond email/password, subscription model, free limits, and trial design.","Health data integrations required for MVP versus post-launch.","Image retention default and whether users may opt in separately to model improvement.","Clinical, dietitian, and certified-trainer review process and launch sign-off."]: bullet(x)

doc.add_heading('18. Recommended next product work', level=1)
flow_num=None
for x in ["Interview 8–12 target users across the three segments; test appetite for calorie ranges, correction flow, and coaching tone.","Prototype the three critical screens: meal capture/result, Today, and adaptive workout.","Create a 100–200 meal benchmark with expert-labeled ranges to compare model quality and clarification strategies.","Define the safety policy and obtain dietitian and certified-trainer review before engineering recommendation logic.","Choose the initial segment, then convert this PRD into a prioritized backlog and technical architecture."]:
    p=numbered(x)
    if flow_num is None: flow_num=restart_numbering(p)
    else: continue_numbering(p,flow_num)

doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run("END OF V0.1 • FOR DISCOVERY AND ALIGNMENT"),8.5,bold=True,color=GRAY)
doc.save(OUT)
print(OUT)
